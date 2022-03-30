import docx
import csv

#pathToTestData = 'C:\\Users\\Oakley\\Downloads\\Example Spreadsheet - Sheet1.csv'
pathToTestData = '/Users/oakleybaron/Downloads/Example Spreadsheet - Sheet1.csv'
# Open csv and read in data

with open(pathToTestData) as csv_file:
    csv_reader = csv.DictReader(csv_file)
    line_count = 1
    donorInfo = []
    
    for row in csv_reader:
        
        list = [(k, v) for k, v in row.items()] 
        name = list[0][1]
        address = list[1][1]
        email = list[2][1]
        phoneNumber = list[3][1]
        date = list[4][1]

        print(name)
        #donorInfo.insert(line_count, row)

        ## Make Word doc for each donor
        # Open template document
        doc = docx.Document()
        
        doc.add_paragraph('Hello, '+ str(name) +'!')
        doc.add_paragraph('Thank you for your donation in '+str(date)+'! Please confirm the following personal information:')
        doc.add_paragraph('Name: '+ name)
        doc.add_paragraph('Address: '+address)
        doc.add_paragraph('Email: '+email)
        doc.add_paragraph('Phone Number: '+phoneNumber)


        # Save new document with data and donor name in title
        doc.save('/Users/oakleybaron/Downloads/'+name+'.docx')

        line_count += 1
